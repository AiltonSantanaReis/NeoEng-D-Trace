from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "NEOENG_EDITOR_COMPOSICAO_2D_NORMATIVO_2026-08-27.md"
OUTPUT = ROOT / "docs" / "NEOENG_EDITOR_COMPOSICAO_2D_NORMATIVO_2026-08-27.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "C9D2DE"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn("w:" + tag))
        if node is None:
            node = OxmlElement("w:" + tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_bottom_border(paragraph, color=BLUE, size="16", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def set_style_font(style, name="Calibri", size=11, color=INK, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.25

    title = doc.styles.add_style("Normative Title", 1)
    set_style_font(title, size=23, color=INK, bold=True)
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.05
    title.paragraph_format.keep_with_next = True

    meta = doc.styles.add_style("Masthead Metadata", 1)
    set_style_font(meta, size=10, color=MUTED)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.10

    table_text = doc.styles.add_style("Table Text", 1)
    set_style_font(table_text, size=9, color=INK)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.10

    code_style = doc.styles.add_style("Reference Text", 1)
    set_style_font(code_style, name="Consolas", size=9, color=INK)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.10


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def create_definition(abstract_id, num_id, fmt, text, left, hanging, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)
        return num_id

    bullet_id = create_definition(
        next_abs, next_num, "bullet", "•", 540, 270, "Calibri"
    )
    decimal_id = create_definition(
        next_abs + 1, next_num + 1, "decimal", "%1.", 540, 270
    )
    return bullet_id, decimal_id


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def split_inline(text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = []
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            parts.append((text[cursor : match.start()], False, False))
        token = match.group(0)
        if token.startswith("**"):
            parts.append((token[2:-2], True, False))
        else:
            parts.append((token[1:-1], False, True))
        cursor = match.end()
    if cursor < len(text):
        parts.append((text[cursor:], False, False))
    return parts or [(text, False, False)]


def add_inline(paragraph, text, size=11, color=INK):
    for part, bold, italic in split_inline(text):
        run = paragraph.add_run(part)
        set_run_font(
            run,
            size=size,
            color=color,
            bold=bold if bold else None,
            italic=italic if italic else None,
        )


def add_text_paragraph(doc, text, style="Normal", size=11, color=INK):
    paragraph = doc.add_paragraph(style=style)
    add_inline(paragraph, text, size=size, color=color)
    return paragraph


def clean_table_row(line):
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_separator_row(cells):
    return bool(cells) and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells
    )


def table_widths(count):
    if count == 1:
        return [9360]
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [1800, 3000, 4560]
    if count == 4:
        return [1500, 1900, 2500, 3460]
    total = CONTENT_DXA
    base = total // count
    widths = [base] * count
    widths[-1] += total - sum(widths)
    return widths


def add_table(doc, rows):
    if not rows:
        return
    count = max(len(row) for row in rows)
    normalized = [row + [""] * (count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=count)
    set_table_geometry(table, table_widths(count))
    set_table_borders(table)
    for row_idx, row in enumerate(normalized):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = "Table Text"
            add_inline(p, value, size=9, color=INK)
            if row_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
            else:
                set_cell_shading(cell, WHITE)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D6DEE8", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.03)
    add_inline(p, text, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metadata_line(doc, text):
    paragraph = doc.add_paragraph(style="Masthead Metadata")
    add_inline(paragraph, text, size=10, color=MUTED)
    return paragraph


def configure_page_furniture(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    add_inline(
        hp, "NeoEng-D-Trace  |  Documento normativo de produto", size=8.5, color=MUTED
    )

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_inline(fp, "Artefato local de engenharia  |  Página ", size=8.5, color=MUTED)
    add_page_field(fp)


def render_markdown(doc, lines):
    bullet_id, decimal_id = add_numbering(doc)
    first_heading = True
    metadata_mode = False
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "---":
            if metadata_mode:
                rule = doc.add_paragraph()
                rule.paragraph_format.space_before = Pt(3)
                rule.paragraph_format.space_after = Pt(8)
                add_bottom_border(rule)
                metadata_mode = False
            else:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            style = "Normative Title" if first_heading else "Heading 1"
            p = add_text_paragraph(
                doc,
                text,
                style=style,
                size=23 if first_heading else 16,
                color=INK if first_heading else BLUE,
            )
            if first_heading:
                first_heading = False
                metadata_mode = True
            i += 1
            continue
        if line.startswith("## "):
            metadata_mode = False
            add_text_paragraph(
                doc, line[3:].strip(), style="Heading 1", size=16, color=BLUE
            )
            i += 1
            continue
        if line.startswith("### "):
            metadata_mode = False
            add_text_paragraph(
                doc, line[4:].strip(), style="Heading 2", size=13, color=BLUE
            )
            i += 1
            continue
        if line.startswith("#### "):
            metadata_mode = False
            add_text_paragraph(
                doc, line[5:].strip(), style="Heading 3", size=12, color=DARK_BLUE
            )
            i += 1
            continue
        if line.startswith(">"):
            callout_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                callout_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_callout(doc, " ".join(callout_lines))
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(clean_table_row(lines[i]))
                i += 1
            if len(table_lines) >= 2 and is_separator_row(table_lines[1]):
                table_lines.pop(1)
            add_table(doc, table_lines)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            set_numbering(p, bullet_id)
            add_inline(p, line[2:].strip(), size=11, color=INK)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            body = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="Normal")
            set_numbering(p, decimal_id)
            add_inline(p, body, size=11, color=INK)
            i += 1
            continue
        if metadata_mode and line.startswith("**"):
            add_metadata_line(doc, line)
            i += 1
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if (
                not candidate
                or candidate == "---"
                or candidate.startswith(("#", ">", "|", "- "))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        add_text_paragraph(
            doc, " ".join(paragraph_lines), style="Normal", size=11, color=INK
        )


def main():
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_page_furniture(doc)
    doc.core_properties.title = (
        "NeoEng-D-Trace — Documento Normativo de Consolidação do "
        "Editor Profissional de Composição 2D Baseado em Objetos"
    )
    doc.core_properties.subject = (
        "Plano obrigatório, imutáveis, gates, evidências e aceite formal"
    )
    doc.core_properties.author = "Equipe de desenvolvimento NeoEng-D-Trace"
    doc.core_properties.comments = (
        "Gerado a partir da fonte Markdown normativa versionável."
    )
    render_markdown(doc, source_text.splitlines())
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
