from pathlib import Path
import re
from docx import Document
from docx.oxml.ns import qn

path = Path("docs/NEOENG_EDITOR_COMPOSICAO_2D_NORMATIVO_2026-08-27.docx")
doc = Document(path)
paragraph_text = [p.text for p in doc.paragraphs]
cell_text = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
all_text = "\n".join(paragraph_text + cell_text)
section = doc.sections[0]

print(f"DOCX_BYTES={path.stat().st_size}")
print(f"READABLE={int(bool(doc.core_properties.title))}")
print(f"SOURCE_SHA_PRESENT={int('7df73f21f5a609196e6cadac85065c89a989ecb4' in all_text)}")
forbidden_placeholder_count = len(re.findall(r"\b(?:TODO|TBD|PLACEHOLDER)\b", all_text, re.I))
print(f"FORBIDDEN_PLACEHOLDER_COUNT={forbidden_placeholder_count}")
print(f"PARAGRAPHS={len(doc.paragraphs)}")
print(f"TABLES={len(doc.tables)}")
print(f"HEADINGS_H1={sum(p.style.name == 'Heading 1' for p in doc.paragraphs)}")
print(f"HEADINGS_H2={sum(p.style.name == 'Heading 2' for p in doc.paragraphs)}")
print(f"NUMBERED_PARAGRAPHS={sum(p._p.pPr is not None and p._p.pPr.numPr is not None for p in doc.paragraphs)}")
print(f"PAGE_DXA={section.page_width.twips},{section.page_height.twips}")
print(f"MARGINS_DXA={section.top_margin.twips},{section.right_margin.twips},{section.bottom_margin.twips},{section.left_margin.twips}")
print(f"HEADER_FOOTER_DXA={section.header_distance.twips},{section.footer_distance.twips}")
print("TABLE_WIDTHS_DXA=" + ",".join(str(sum(int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.gridCol_lst)) for table in doc.tables))
indents = []
for table in doc.tables:
    node = table._tbl.tblPr.find(qn("w:tblInd"))
    if node is not None:
        indents.append(str(int(node.get(qn("w:w")))))
print("TABLE_INDENTS_DXA=" + ",".join(indents))
print(f"NONEMPTY_TABLE_CELLS={sum(bool(cell.text.strip()) for cell in [cell for table in doc.tables for row in table.rows for cell in row.cells])}")
